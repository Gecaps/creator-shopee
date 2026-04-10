#!/usr/bin/env python3
"""Gera o roteiro das aulas em formato Word (.docx) visual e profissional."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# ─── Estilos ───
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.3

SHOPEE_ORANGE = RGBColor(0xEE, 0x4D, 0x2D)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x66, 0x66, 0x66)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def add_heading_styled(text, level=1, color=DARK):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h

def add_para(text, bold=False, color=DARK, size=11, align=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if align:
        p.alignment = align
    return p

def add_script_block(text):
    """Adiciona bloco de roteiro com fundo visual."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.font.name = 'Calibri'
    run.italic = True
    return p

def add_task_box(text):
    """Adiciona tarefa destacada."""
    p = doc.add_paragraph()
    run = p.add_run('TAREFA: ')
    run.bold = True
    run.font.color.rgb = SHOPEE_ORANGE
    run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    return p

def add_screen_note(text):
    """Nota do que mostrar na tela."""
    p = doc.add_paragraph()
    run = p.add_run('TELA: ')
    run.bold = True
    run.font.color.rgb = GREEN
    run.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.italic = True
    return p

def add_divider():
    p = doc.add_paragraph()
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    run.font.size = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ─── CAPA ───
for _ in range(6):
    doc.add_paragraph()

add_para('DESAFIO CREATOR SHOPEE', bold=True, color=SHOPEE_ORANGE, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Roteiro Completo das 15 Aulas', bold=True, color=DARK, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('Guia de gravação para a apresentadora', color=GRAY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Cada aula tem: roteiro do que falar, o que mostrar na tela e tarefa prática.', color=GRAY, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('Creator Brasil — 2026', color=GRAY, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

doc.add_page_break()

# ─── CRONOGRAMA ───
add_heading_styled('Cronograma de Gravação', 1, SHOPEE_ORANGE)
add_para('Reserve 4-5 horas por dia de gravação (com pausas e repetições).', color=GRAY, size=10)

table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Dia', 'Módulo', 'Aulas', 'Duração']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

data = [
    ['Dia 1', 'Módulo 0 — Comece por Aqui', 'Aulas 1-3', '~20 min'],
    ['Dia 2', 'Módulo 1 — Preparação', 'Aulas 4-6', '~30 min'],
    ['Dia 3', 'Módulo 2 — Primeiras Vendas', 'Aulas 7-9', '~30 min'],
    ['Dia 4', 'Módulo 3 — Acelerando', 'Aulas 10-12', '~35 min'],
    ['Dia 5', 'Módulo 4 — Escalando', 'Aulas 13-15', '~30 min'],
]
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val
        for p in table.rows[i+1].cells[j].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph()
add_para('Total: ~2h25 de conteúdo final', bold=True, size=12)

doc.add_paragraph()
add_heading_styled('Setup de Gravação', 2, DARK)
items = [
    'Câmera: celular na horizontal, 1080p no mínimo',
    'Áudio: microfone de lapela (R$30-50) — áudio ruim mata o vídeo',
    'Iluminação: de frente pra janela ou ring light',
    'Fundo: parede limpa ou cenário com produtos Shopee',
    'Roupa: mesma roupa dentro do mesmo módulo',
    'Legendas: adicionar depois (CapCut grátis faz automático)',
    'Tela do celular: gravar separado usando gravação de tela nativa',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(10)

doc.add_page_break()

# ─── AULAS ───

aulas = [
    # Módulo 0
    {
        'modulo': 'MÓDULO 0 — COMECE POR AQUI',
        'num': 1,
        'titulo': 'Bem-vindo ao Desafio',
        'duracao': '3-5 min',
        'objetivo': 'Acolher, gerar empolgação e criar compromisso.',
        'roteiro': '''E aí! Seja muito bem-vinda ao Desafio Creator Shopee.

Se você tá aqui, é porque decidiu que quer ganhar dinheiro de verdade, no seu tempo, do seu celular. E você tomou a decisão certa.

Nos próximos 30 dias, eu vou te pegar pela mão e te levar do zero absoluto até sua primeira comissão como afiliada da Shopee.

Deixa eu te explicar como vai funcionar:

O desafio tem 5 módulos e 15 aulas. Cada aula tem uma tarefa prática — você assiste, faz, e pronto. Não precisa de computador, não precisa saber de tecnologia. Se você sabe mandar um zap, você sabe fazer isso.

Mas eu preciso combinar uma coisa com você: isso aqui não é pra assistir e guardar. É pra FAZER. A mágica tá na prática, não na teoria.

Então fecha o Instagram, silencia o WhatsApp, e vamos começar.

Na próxima aula eu te mostro como usar a plataforma. Te vejo lá.''',
        'tarefa': 'Nenhuma — só se comprometer a fazer o desafio.',
        'tela': 'Só ela falando pra câmera. Texto aparecendo: "30 dias", "15 aulas", "1 tarefa por dia".',
    },
    {
        'num': 2,
        'titulo': 'Como usar a plataforma',
        'duracao': '3-4 min',
        'objetivo': 'Ensinar a navegar no MemberKit sem se perder.',
        'roteiro': '''Antes de começar o conteúdo, deixa eu te mostrar como a plataforma funciona pra você não se perder.

Aqui na esquerda você tem os módulos. Hoje a gente tá no Módulo 0. Os outros vão liberando conforme você avança.

Cada módulo tem 3 aulas. Cada aula tem uma tarefa. Faz a tarefa antes de ir pra próxima, combinado?

Se tiver dúvida em qualquer momento, vai no grupo do WhatsApp que eu vou te mostrar na próxima aula.

Pronto, é só isso. Simples. Agora vamos pro que interessa.''',
        'tarefa': 'Navegue pela plataforma e veja os módulos.',
        'tela': 'Gravação de tela do MemberKit navegando pelos módulos. Setas nos botões importantes.',
    },
    {
        'num': 3,
        'titulo': 'Entre no grupo do WhatsApp',
        'duracao': '2-3 min',
        'objetivo': 'Levar pro grupo de suporte.',
        'roteiro': '''Essa aula é rápida mas é importante.

Eu quero que você entre no nosso grupo do WhatsApp agora. Lá é onde a gente tira dúvida, compartilha resultado e se ajuda.

Porque uma coisa que eu aprendi: quem faz sozinho, desiste. Quem faz com gente junto, continua.

O link tá aqui embaixo da aula. Clica, entra no grupo, e manda um "cheguei!" lá.

Te vejo no grupo e na próxima aula, onde a gente começa de verdade.''',
        'tarefa': 'Entre no grupo do WhatsApp e mande "cheguei!".',
        'tela': 'O link do grupo. Screenshot do grupo com mensagens de boas-vindas.',
    },
    # Módulo 1
    {
        'modulo': 'MÓDULO 1 — PREPARAÇÃO (Dias 1-3)',
        'num': 4,
        'titulo': 'Como funciona o programa de afiliados da Shopee',
        'duracao': '7-10 min',
        'objetivo': 'Explicar o modelo de negócio de forma simples e concreta.',
        'roteiro': '''Agora vamos pro conteúdo de verdade. Antes de criar qualquer coisa, você precisa entender COMO você vai ganhar dinheiro. É simples:

A Shopee tem milhões de produtos. Eles querem que mais gente compre esses produtos. Então criaram o Programa de Afiliados: você divulga um produto, alguém compra pelo seu link, e a Shopee te paga uma comissão.

Você NÃO compra o produto. Você NÃO estoca nada. Você NÃO entrega nada. A Shopee faz tudo isso. Você só indica.

Agora, quanto você ganha? Depende da categoria:
- Moda e acessórios: 8% a 10%
- Beleza: 8% a 12%
- Eletrônicos: 5% a 7%
- Em campanhas especiais, pode chegar a 20-30%

Exemplo prático: se você indicar um kit de maquiagem de R$150 com 10% de comissão, você ganha R$15. Se 20 pessoas comprarem no mês, são R$300. Se forem 50 pessoas, R$750.

O link de afiliado tem um cookie de 7 dias — se a pessoa clicar hoje e comprar daqui 5 dias, a comissão ainda é sua.

O pagamento é via ShopeePay em 30 a 60 dias. Mínimo pra sacar: R$30.''',
        'tarefa': 'Entenda o modelo. Não precisa fazer nada ainda — só entenda que quando alguém compra pelo seu link, você ganha dinheiro.',
        'tela': 'Tabela de comissões por categoria. Exemplo prático com números. Diagrama: "Você indica → Pessoa compra → Shopee paga você".',
    },
    {
        'num': 5,
        'titulo': 'Criando sua conta de Creator',
        'duracao': '8-10 min',
        'objetivo': 'Passo a passo de criar a conta no programa de afiliados.',
        'roteiro': '''Agora é hora de colocar a mão na massa. Vamos criar sua conta de creator na Shopee. Pega o celular e faz junto comigo.

Passo 1: Abre o app da Shopee. Se não tem, baixa na loja.

Passo 2: Faz login na sua conta. Se não tem, cria uma — é grátis.

Passo 3: Vai em "Eu" lá embaixo → rola até achar "Central do Creator" ou "Programa de Afiliados".

Passo 4: Clica em "Quero participar". Vai pedir: nome completo, CPF, suas redes sociais.

Passo 5: Confirma e pronto. Sua conta de afiliado tá ativa.

Detalhe importante: pra algumas funções, a Shopee pede seguidores. Mas pra gerar links de afiliado básicos, não precisa de nada.''',
        'tarefa': 'CRIE SUA CONTA agora. Tira um print e manda no grupo com "conta criada!".',
        'tela': 'Gravação de tela mostrando CADA CLIQUE no app da Shopee. Setas indicando onde clicar.',
    },
    {
        'num': 6,
        'titulo': 'Escolhendo seus primeiros produtos',
        'duracao': '8-10 min',
        'objetivo': 'Ensinar a escolher produtos com boa comissão e demanda real.',
        'roteiro': '''Conta criada? Ótimo. Agora vem a parte mais importante: escolher os produtos certos.

Critérios pra um bom produto:
1. Comissão de 8% pra cima
2. Demanda real (as pessoas já querem comprar)
3. Avaliação de 4 estrelas ou mais
4. Preço entre R$50 e R$300

Vai no app → Central do Creator → "Produtos pra divulgar".

As melhores categorias pra começar:
- Beleza e skincare — comissão alta, mulheres compram muito
- Moda feminina — volume enorme
- Casa e decoração — ticket médio bom
- Acessórios de celular — todo mundo precisa

Evite no começo: eletrônicos caros (comissão baixa), produtos sem avaliação, produtos com frete caro.''',
        'tarefa': 'Escolha 5 produtos pra divulgar. Salva nos favoritos. Manda no grupo quais escolheu.',
        'tela': 'Navegação no marketplace de afiliados. Exemplos de produtos bons vs ruins. Filtros por comissão.',
    },
    # Módulo 2
    {
        'modulo': 'MÓDULO 2 — PRIMEIRAS VENDAS (Dias 4-10)',
        'num': 7,
        'titulo': 'Como criar links de afiliado',
        'duracao': '5-7 min',
        'objetivo': 'Gerar os primeiros links prontos pra divulgar.',
        'roteiro': '''Agora vamos criar seus links de afiliado. É isso que conecta a venda a VOCÊ.

Vai no produto que você escolheu. Clica em "Gerar link de afiliado" ou "Compartilhar como afiliado".

A Shopee gera um link especial — esse link é SÓ seu. Toda compra por esse link nos próximos 7 dias é comissão sua.

Duas formas de compartilhar:
1. O link direto — copia e cola
2. A imagem do produto — a Shopee gera uma imagem bonita com preço

Dica: salva seus links num bloco de notas no celular. Cria uma nota "Meus links Shopee" e vai colando ali.''',
        'tarefa': 'Gere o link dos 5 produtos. Salva todos no bloco de notas.',
        'tela': 'Geração do link passo a passo no app. Como copiar. Como salvar.',
    },
    {
        'num': 8,
        'titulo': 'Divulgando no WhatsApp',
        'duracao': '10-12 min',
        'objetivo': 'Primeira estratégia de divulgação — a mais acessível. AULA MAIS IMPORTANTE.',
        'roteiro': '''Essa é a aula mais importante do curso. É aqui que o dinheiro começa a entrar.

O WhatsApp é sua ferramenta mais poderosa. Todo mundo que você conhece tá lá, e as pessoas confiam em indicação de gente que conhecem.

Mas ATENÇÃO: não saia spamando. A estratégia certa tem 3 frentes:

1. STATUS DO WHATSAPP (a mina de ouro)
- Posta o produto como indicação natural: "Gente, achei esse kit de skincare por R$79 na Shopee e tá com 4.8 estrelas. Link nos comentários."
- 2-3 status por dia, em horários diferentes
- Alterna: achados, promoções, reviews
- NÃO posta 10 produtos seguidos, senão vira catálogo

2. GRUPOS
- Vira a pessoa que acha promoção: "Achei esse produto com 40% off, quem quiser o link é esse"
- Crie um grupo SEU: "Ofertas e Promoções" — posta 2-3 ofertas/dia

3. MENSAGEM DIRETA (com cuidado)
- Se alguém postar que tá procurando algo, manda: "Vi que você tá procurando X, achei esse ótimo na Shopee"
- Indicação natural, não spam''',
        'tarefa': 'Hoje: 1) Poste 2 produtos no status. 2) Mande 1 oferta num grupo. 3) Mande no grupo do curso: "postei meu primeiro status!".',
        'tela': 'Exemplo de status bem feito. Exemplo de mensagem em grupo. Horários recomendados.',
    },
    {
        'num': 9,
        'titulo': 'Divulgando no Instagram e TikTok',
        'duracao': '8-10 min',
        'objetivo': 'Expandir para redes sociais (OPCIONAL).',
        'roteiro': '''Se o WhatsApp é o feijão com arroz, o Instagram e TikTok são o churrasco. Mais trabalho, mais resultado.

IMPORTANTE: essa aula é OPCIONAL. O WhatsApp já funciona. Mas se quiser acelerar:

INSTAGRAM — crie uma conta de "achados":
- Nome tipo @achadosdashopee
- Posta fotos dos produtos com preço e link na bio
- Stories com reviews, reels de 15-30s mostrando achados
- NÃO precisa aparecer — pode ser só foto + texto + preço

TIKTOK — onde o jogo tá mudando:
- Com TikTok Shop, marca o produto direto no vídeo
- Formatos: "achados que você precisa", "testei esse produto", "desconto hoje"
- Segredo: vídeo feito no celular, sem edição, performa MELHOR''',
        'tarefa': 'Se quiser ir além: escolha UMA rede e faça 1 post hoje. Se não quiser, continue no WhatsApp.',
        'tela': 'Exemplos de perfis de achados. Exemplo de Reel/TikTok. Como colocar link na bio.',
    },
    # Módulo 3
    {
        'modulo': 'MÓDULO 3 — ACELERANDO (Dias 11-20)',
        'num': 10,
        'titulo': 'Criando conteúdo que vende',
        'duracao': '10-12 min',
        'objetivo': 'Ensinar a diferença entre postar produto e criar conteúdo que converte.',
        'roteiro': '''Você já tá divulgando. Agora vamos MELHORAR a qualidade.

Regra de ouro: NÃO VENDA O PRODUTO, VENDA O RESULTADO.

Ninguém quer comprar sérum de vitamina C. As pessoas querem pele bonita.

Em vez de: "Sérum vitamina C por R$39"
Faz: "Minha pele tava manchada e em 2 semanas com esse sérum melhorou demais. R$39 na Shopee."

5 tipos de conteúdo que mais vendem:
1. Review pessoal — "comprei e testei, olha o resultado"
2. Antes e depois — demais pra beleza, organização
3. Comparação — "testei o de R$30 e o de R$150"
4. Achado do dia — "não acredito nesse preço"
5. Urgência — "promoção acaba hoje, corre"''',
        'tarefa': 'Pega 1 produto e cria conteúdo usando uma das 5 fórmulas. Posta.',
        'tela': 'Exemplos lado a lado (post ruim vs bom). As 5 fórmulas escritas na tela.',
    },
    {
        'num': 11,
        'titulo': 'Encontrando produtos com alta comissão',
        'duracao': '8-10 min',
        'objetivo': 'Ir além dos produtos básicos.',
        'roteiro': '''A conta é simples:
- Produto de R$30 com 5% = R$1,50/venda → precisa de 667 vendas pra R$1.000
- Produto de R$200 com 12% = R$24/venda → precisa de 42 vendas pra R$1.000

Regra: PRIORIZE TICKET ALTO + COMISSÃO ALTA.

Onde achar:
1. Campanhas da Shopee (7.7, 8.8, 11.11) — comissões sobem
2. Parceiros selecionados — aba de ofertas exclusivas, até 30%
3. Produtos em alta — trending da Shopee
4. Nichos de alto ticket: eletrodomésticos portáteis, skincare premium, suplementos, cama/mesa/banho''',
        'tarefa': 'Encontre 3 produtos novos com comissão de 10%+ e ticket acima de R$100.',
        'tela': 'Painel com filtro de comissão. Comparação matemática. Exemplos de alto ticket.',
    },
    {
        'num': 12,
        'titulo': 'Estratégias para aumentar cliques',
        'duracao': '8-10 min',
        'objetivo': 'Otimizar a divulgação.',
        'roteiro': '''6 estratégias pra mais cliques:

1. HORÁRIOS CERTOS: 7h-9h, 12h-13h, 19h-22h (melhor horário)

2. ESCASSEZ: "só tem 3 em estoque", "cupom até meia-noite"

3. PROVA SOCIAL: "já vendeu 5.000 unidades", "nota 4.9"

4. PREÇO ÂNCORA: "De R$199 por R$89 — mais da metade off"

5. FREQUÊNCIA: poste TODO DIA. Consistência separa quem ganha R$200 de quem ganha R$3.000

6. PEÇA PRA COMPARTILHAR: "manda pra alguém que ia gostar"''',
        'tarefa': 'Poste HOJE entre 19h-21h usando pelo menos 2 dessas estratégias.',
        'tela': 'Gráfico de horários. Exemplos de posts com gatilhos.',
    },
    # Módulo 4
    {
        'modulo': 'MÓDULO 4 — ESCALANDO (Dias 21-30)',
        'num': 13,
        'titulo': 'Analisando seus resultados',
        'duracao': '8-10 min',
        'objetivo': 'Ensinar a ler o painel e entender o que funciona.',
        'roteiro': '''Dia 21. Agora a gente olha pros NÚMEROS. O que não se mede, não melhora.

No painel da Shopee você vê: cliques, vendas, comissão, produtos que mais venderam.

TAXA DE CONVERSÃO = vendas ÷ cliques × 100
- Abaixo de 1%: troque o produto
- Acima de 5%: duplique o esforço nele

Identifique seus produtos campeões e foque neles. Pare de divulgar o que ninguém quer.''',
        'tarefa': 'Abra o painel, anote seus números, identifique os 3 melhores produtos.',
        'tela': 'Painel real com métricas. Como calcular conversão.',
    },
    {
        'num': 14,
        'titulo': 'Escalando de R$500 para R$3.000',
        'duracao': '10-12 min',
        'objetivo': 'Multiplicar os ganhos.',
        'roteiro': '''Se você seguiu o desafio, já tá entre R$200 e R$500. Como ir pra R$3.000?

MULTIPLICAÇÃO, não revolução. Faça MAIS do que já funciona.

5 estratégias:

1. MAIS CANAIS: WhatsApp (R$500) + Instagram (R$1.500) + TikTok (R$3.000+)

2. MAIS PRODUTOS: de 5 pra 15-20 produtos

3. DATAS COMERCIAIS: 5.5, 7.7, 11.11, Black Friday — comissões de 20-30%

4. CONTEÚDO EVERGREEN: posts que vendem por semanas

5. SUA AUDIÊNCIA: grupo de ofertas no WhatsApp/Telegram — meta: 500 pessoas''',
        'tarefa': 'Escolha UMA estratégia e comece a implementar hoje.',
        'tela': 'Conta da multiplicação. Calendário de campanhas. Exemplos de conteúdo evergreen.',
    },
    {
        'num': 15,
        'titulo': 'Próximos passos e plano de ação',
        'duracao': '5-7 min',
        'objetivo': 'Fechar o desafio, celebrar e dar direção.',
        'roteiro': '''Você completou o Desafio Creator Shopee. Parabéns.

A maioria nunca termina um curso online. Você terminou. Já tá na frente de 90%.

O que fazer agora:
1. Continue postando todo dia
2. Expanda seus canais
3. Fique de olho nas campanhas
4. Continue no grupo
5. Pense grande — os top afiliados ganham R$20.000/mês

Lembra: R$300/mês já paga a conta de luz. R$1.000 já muda sua vida. R$3.000 já é renda real. Não menospreza os pequenos resultados.

Se fez sentido, compartilha com alguém que precisa de renda extra.

Obrigada por ter confiado no nosso método. Te vejo no grupo. Bora que bora!''',
        'tarefa': 'Monte seu plano de ação pro próximo mês. Mande no grupo.',
        'tela': 'Checklist do que conquistou. Os 5 próximos passos.',
    },
]

current_modulo = None
for aula in aulas:
    if 'modulo' in aula and aula['modulo'] != current_modulo:
        current_modulo = aula['modulo']
        doc.add_page_break()
        add_heading_styled(current_modulo, 1, SHOPEE_ORANGE)
        add_divider()

    doc.add_paragraph()
    add_heading_styled(f"Aula {aula['num']}: {aula['titulo']}", 2, DARK)

    # Info bar
    p = doc.add_paragraph()
    run = p.add_run(f"Duração: {aula['duracao']}  |  ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = SHOPEE_ORANGE
    run = p.add_run(aula['objetivo'])
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.italic = True

    # Roteiro
    add_para('O QUE FALAR:', bold=True, size=10, color=SHOPEE_ORANGE)
    for paragraph in aula['roteiro'].split('\n\n'):
        if paragraph.strip():
            add_script_block(paragraph.strip())

    doc.add_paragraph()
    add_task_box(aula['tarefa'])
    add_screen_note(aula['tela'])
    add_divider()

# ─── NOTAS FINAIS ───
doc.add_page_break()
add_heading_styled('Notas de Produção', 1, SHOPEE_ORANGE)

add_heading_styled('Tom de Voz', 2, DARK)
items = [
    'Conversa de amiga, não palestra',
    'Usa "você", não "vocês"',
    'Explica como se fosse pra alguém que nunca fez isso',
    'Empolgada mas natural — pode rir, errar, recomeçar',
    'Autenticidade > perfeição',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('Estrutura de Cada Aula', 2, DARK)
items = [
    'Abertura (10s): "Nessa aula você vai aprender X"',
    'Conteúdo: ensina o que prometeu',
    'Tarefa: uma ação específica pra fazer AGORA',
    'Gancho: "Na próxima aula, a gente vai..."',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('O que NÃO Fazer', 2, DARK)
items = [
    'Não prometa valores específicos ("você VAI ganhar R$3.000")',
    'Diga "o método permite chegar a" ou "nossos alunos chegam a"',
    'Não seja condescendente — o público é adulto',
    'Não enrole — se pode explicar em 2 min, não leva 10',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ─── SALVAR ───
output_path = '/Users/work/gecaps/PROJETOS/CREATOR SHOPEE/copy/Roteiro-Aulas-Creator-Shopee.docx'
doc.save(output_path)
print(f'Salvo em: {output_path}')
